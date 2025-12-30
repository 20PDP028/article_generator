from docx import Document


def export_to_word(
    article,
    intro_count,
    discussion_count,
    conclusion_count,
    references,
    filename="final_article.docx"
):
    doc = Document()

    # -------- TITLE --------
    doc.add_heading("Rewritten Review Article Draft", level=1)

    index = 0

    # -------- INTRODUCTION --------
    doc.add_heading("Introduction", level=2)
    for _ in range(intro_count):
        paragraph, _ = article[index]
        doc.add_paragraph(" ".join(paragraph))
        index += 1

    # -------- DISCUSSION --------
    doc.add_heading("Discussion", level=2)
    for _ in range(discussion_count):
        paragraph, _ = article[index]
        doc.add_paragraph(" ".join(paragraph))
        index += 1

    # -------- CONCLUSION --------
    doc.add_heading("Conclusion", level=2)
    for _ in range(conclusion_count):
        paragraph, _ = article[index]
        doc.add_paragraph(" ".join(paragraph))
        index += 1

    # -------- REFERENCES --------
    if references:
        doc.add_heading("References", level=2)
        for i, ref in enumerate(references, 1):
            doc.add_paragraph(f"{i}. {ref}")

    doc.save(filename)
